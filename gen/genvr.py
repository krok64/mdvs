# -*- coding: utf-8 -*-

from django.http import HttpResponse
from django.shortcuts import render, redirect
from math import pi
from tempfile import NamedTemporaryFile
from gen.models import *
from docx import *
from . biur import GetBiur

#Формат вывода
AreaFormat="{:.2f}"
MassaFormat="{:.1f}"
ShtukFormat="{:.0f}"
MetrFormat="{:.1f}"

#Перевод строки
CR="$$$"

def TipMarka(detal):
    # выбираем что писать в графу Тип/марка в СО
    if detal.s_name:
        tip_marka=detal.s_name.encode('utf-8')
    elif detal.tu:
        tip_marka=detal.tu.name.encode('utf-8')
        if detal.tu2:
            tip_marka=tip_marka+CR+detal.tu2.name.encode('utf-8')
    else:
        tip_marka=" ".encode('utf-8')

    #выбираем что писать в графу Завод-изготовитель
    if detal.tu and detal.tu.zavod.name:
        zavod_izg=detal.tu.zavod.name.encode('utf-8')
        if zavod_izg=='Нет':
            zavod_izg=''
    else:
        zavod_izg=" ".encode('utf-8')

    return (tip_marka, zavod_izg)

def gen_SO(request, spec_id):
    s=Spec.objects.get(pk=spec_id)
    p=ItemList.objects.filter(spec=s)

    #тут хранится новая спецификация
    l=[]

#    l.append(("ОБОРУДОВАНИЕ, ПОСТАВЛЯЕМОЕ ЗАКАЗЧИКОМ",))
#Вся запорно-регулирующая арматура: задвижки, краны, вентили, клапаны, кроме до Ду200 с механическим приводом
#Фильтры-осушители
#камеры приема/запуска
#емкости
#колодцы
#муфты электроизолирующие

#выводим все вместе
    #врезка под давлением
    pc=p.filter(item__type__name='Врезка под давлением').order_by('item__name')
    for k in pc:
        tip_marka, zavod_izg=TipMarka(k.item)
        l.append((k.item.name.encode('utf-8'), tip_marka,'',
            zavod_izg, "шт", sf(k.numAll), kf(k.item.m)))

    #краны
    pc=p.filter(item__type__name='Кран').order_by('-item__dy')
    for k in pc:
        tip_marka, zavod_izg=TipMarka(k.item)
        l.append((k.item.name.encode('utf-8'), tip_marka,'',
            zavod_izg, "шт", sf(k.numAll), kf(k.item.m)))

    #трубы
    pc=p.filter(item__type__name='Труба').order_by('-item__dy')
    for k in pc:
        tip_marka, zavod_izg=TipMarka(k.item)
        #вычисляем к-т отхода труб (<500 - 1.01, 600-1000 - 1.008, >1000 - 1.006)
        if k.item.d < 500:
            k_othoda_trub=1.01
        elif k.item.d>500 and k.item.d<1200:
            k_othoda_trub=1.008
        else:
            k_othoda_trub=1.006
#длину труб умножаем на к-т отхода труб и округляем в большую сторону с точностью 0.1
        l.append((k.item.name.encode('utf-8'), tip_marka,'',
            zavod_izg, "м", mf(k.numAll*k_othoda_trub+0.049), kf(k.item.m),'k=%.3f'% (k_othoda_trub)))

    #Список деталей по которым идет газ (кроме труб)
    GasList=['Тройник','Переход','Отвод','Заглушка','Днище','Фланец']
    pc=p.filter(item__type__name__in=GasList).order_by('item__type__name', '-item__dy')
    for k in pc:
        tip_marka, zavod_izg=TipMarka(k.item)
        comment=''
        detal_name=k.item.name.encode('utf-8')
        #согласно протоколу 35 для деталей по ГОСт указываем доп информацию
        if k.item.tu.name.encode('utf-8') in ('ГОСТ 17375-2001','ГОСТ 17376-2001','ГОСТ 17378-2001','ГОСТ 17379-2001'):
            if k.item.s<=10:
                kcu=34.3
            elif k.item.s<=25:
                kcu=49.0
            else:
                kcu=58.8
            if s.ClimIsp == CLIMAT_NORM:
                temp1=-40
                temp2=-5
            else:
                temp1=-60
                temp2=-20
            ppr=float(s.NumTry.split()[0])
            comment='толщина стенки s рассчитана по СП 36.13330.2012'
            detal_name=detal_name+' (Рраб=%.1fМПа, Рпр=%.1fРраб, KCU не менее %.1f Дж/см2 при t=%dC, KCV не менее 35.0 Дж/см2 при t=%dC)' % (s.Prab, ppr,kcu,temp1,temp2)
        l.append((detal_name, tip_marka,'',
            zavod_izg, "шт", sf(k.numAll), kf(k.item.m), comment))

    #все остальное - монтажные
    pc=p.filter(item__type__name='Монтажные').order_by('item__name')
    for k in pc:
        tip_marka, zavod_izg=TipMarka(k.item)
        l.append((k.item.name.encode('utf-8'), tip_marka,'',
            zavod_izg, "шт", sf(k.numAll), kf(k.item.m)))

    #Считаем БИУРС
    GasList=['Тройник','Переход','Отвод','Заглушка','Днище','Кран', 'Труба']
    pc=p.filter(item__type__name__in=GasList, numPod__gt=0, item__zav_is=False)
    biurs=0
    for k in pc:
        if not k.item.d2:
            d2=0
        else:
            d2=k.item.d2
        if k.item.dy:
            d1=k.item.dy
        else:
            d1=GetDU(k.item.d)
        biurs=biurs+GetBiur(k.item.type.name.encode('utf-8'), d1, GetDU(d2))*k.numPod
    l.append(('Система антикоррозионного покрытия «БИУРС», в том числе:'+CR+
        '- Грунтовка эпоксидная быстроотверждающаяся «Праймер МБ»'+CR+
        '- Мастика битумно-уретановая «БИУР»',
        'ТУ 51-31323949-80-2004'+CR+'ТУ 2225-015-00396558-01'+CR+'ТУ 5.966-11610-99','',
        'ООО «БИУРС»,'+CR+'г. Санкт-Петербург', 'кг'+CR+'кг',
        "%.1f %s %.1f" % (biurs*0.08, CR, biurs*0.92)))

    document=Document("templateSO.docx")
    t1=document.tables[0]
    i=0
    for k in l:
        for j in range(8):
            if len(k)<=j:
                break
            if k[j]:
                t1.cell(i,j+1).text=k[j].decode('utf-8')
        i=i+2
        t1.add_row()
        t1.add_row()
    f=NamedTemporaryFile(delete=False)
    f.close()
    document.save(f.name)
    f=open(f.name,'rb')
    response = HttpResponse(f, content_type='text/html')
    response['Content-Disposition'] = 'attachment; filename="SO.docx"'
    return response

def gen_VR(request, spec_id):
    s=Spec.objects.get(pk=spec_id)
    p=ItemList.objects.filter(spec=s)

    #тут хранится новая ведомость
    l=[]

    #Список деталей по которым идет газ (кроме труб)
    GasList=['Тройник','Переход','Отвод','Заглушка','Днище']

#Список объектов для расчета объема гидроиспытания
    GasFullList=['Тройник','Переход','Отвод','Кран','Труба']

    #Инструкция
    l.append(("Инструкция по доработке текста: Заменить $$$ на ^l, Выбрать колонки Ед.Изм. и Количество и выровнять их по центру, Удалить эту строку",))

    #Описание проекта
    l.append((s.comment.encode('utf-8'),))

    #подземная арматура с заводской изоляцией
    pc=p.filter(item__type__name='Кран',item__ust=UST_POD,item__zav_is=True)
    if pc:
        l.append(("Монтаж подземной арматуры с заводской изоляцией:",))
        for k in pc:
            l.append(('-'+k.item.name.encode('utf-8'), "шт"+CR+"кг", sf(k.numPod)+CR+kf(k.numPod*k.item.m)))

    #монтаж монтажных изделий с заводской изоляцией
    pc=p.filter(item__type__name='Монтажные',item__zav_is=True)
    if pc:
        l.append(("Монтаж с заводской изоляцией:",))
        for k in pc:
            l.append(('-'+k.item.name.encode('utf-8'), "шт"+CR+"кг", sf(k.numAll)+CR+kf(k.numAll*k.item.m)))

    #подземные трубопроводы и фасонные части с изоляцией, группировать по Ду, по убыванию
    #для каждого ду отдельно суммируем общую длину
    flagTitle=False
    for dy in reversed(sorted(DYD.keys())):
        summ=0
        flagTruba=False
        listIdx=0
        for d in DYD[dy]:
            pc=p.filter(item__type__name='Труба', item__d=d, item__zav_is=True,numAll__gt=0)
            if pc:
                flagTruba=True
                if not flagTitle:
                    flagTitle=True
                    l.append(('Монтаж с заводской изоляцией:',))
                    l.append(['трубопроводов Ду '+ str(dy) + ' мм в том числе:','м'])
                else:
                    l.append(['то же, трубопроводов Ду '+ str(dy) + ' мм в том числе:','м'])
                listIdx=len(l)
                for k in pc:
                    if k.item.tu:
                        TuName=k.item.tu.name.encode('utf-8')
                    else:
                        TuName=""
                    l.append(('-'+k.item.name.encode('utf-8')+' '+TuName,
                        'м'+CR+'кг',
                        mf(k.numPod)+CR+kf(k.numPod*k.item.m)))
                    summ+=k.numPod
        for d in DYD[dy]:
            pc=p.filter(item__type__name__in=GasList, item__d=d, item__zav_is=True, numAll__gt=0)
            if pc:
                if not flagTitle:
                    flagTitle=True
                    l.append(('Монтаж с заводской изоляцией:',))
                    l.append(['стальных фасонных частей Ду '+ str(dy) + ' мм в том числе:','м'])
                    listIdx=len(l)
                elif not flagTruba:
                    l.append(['то же, стальных фасонных частей Ду '+ str(dy) + ' мм в том числе:','м'])
                    listIdx=len(l)
                for k in pc:
                    l.append(('-'+k.item.name.encode('utf-8')+' '+k.item.tu.name.encode('utf-8'),'шт'+CR+'кг', sf(k.numPod)+CR+kf(k.numPod*k.item.m)))
                    summ+=k.numPod*k.item.dlina/1000
        if summ>0:
            l[listIdx-1].append(mf(summ))

    #изоляция сварных стыков ТУМами
    pc=p.filter(item__d__gt=0,item__zav_is=True)
    if pc:
        l.append(("Изоляция сварных стыков ТУМами:",))
        DList={}
        for k in pc:
            DList[k.item.d]=1
        for i in sorted(DList.keys()):
            l.append((str(i), "компл.", "XXX"))

    # подземные трубопроводы без изоляции
    # в цикле перебираем все ДУ для труб и для фасонных частей начиная с максимальных
    # и выбираем подземные трубы и фасонину без изоляции для каждого Ду
    flagTitle=False
    IzolPodM2=0.0
    for dy in reversed(sorted(DYD.keys())):
        flagTruba=False
        summ=0.0
        listIdx=0
        for d in DYD[dy]:
            pc=p.filter(item__type__name='Труба', item__d=d, item__zav_is=False,numPod__gt=0)
            if pc:
                flagTruba=True
                if not flagTitle:
                    flagTitle=True
                    l.append(('Монтаж, укладка и изоляция подземных',))
                    l.append(['трубопроводов Ду '+ str(dy) + ' мм в том числе:','м'+CR+'м2'])
                else:
                    l.append(['то же, трубопроводов Ду '+ str(dy) + ' мм в том числе:','м'+CR+'м2'])
                listIdx=len(l)
                for k in pc:
                    if k.item.tu:
                        TuName=k.item.tu.name.encode('utf-8')
                    else:
                        TuName=""
                    l.append(('-'+k.item.name.encode('utf-8')+' '+ TuName,
                        'м'+CR+'кг',
                        mf(k.numPod)+CR+kf(k.numPod*k.item.m)))
                    summ+=k.numPod
        for d in DYD[dy]:
            pc=p.filter(item__type__name__in=GasList, item__d=d, item__zav_is=False, numPod__gt=0)
            if pc:
                if not flagTitle:
                    flagTitle=True
                    l.append(('Монтаж, укладка и изоляция подземных',))
                    l.append(['стальных фасонных частей Ду '+ str(dy) + ' мм в том числе:','м'+CR+'м2'])
                    listIdx=len(l)
                elif not flagTruba:
                    l.append(['то же, стальных фасонных частей Ду '+ str(dy) + ' мм в том числе:','м''м'+CR+'м2'])
                    listIdx=len(l)
                for k in pc:
                    l.append(('-'+k.item.name.encode('utf-8')+' '+k.item.tu.name.encode('utf-8'),
                        'шт'+CR+'кг',
                        sf(k.numPod)+CR+kf(k.numPod*k.item.m)))
                    summ+=k.numPod*k.item.dlina/1000
        if summ>0:
            l[listIdx-1].append(mf(summ)+CR+af(AreaTrub(DYD[dy][0],summ)))
            IzolPodM2=IzolPodM2+AreaTrub(DYD[dy][0],summ)

    #Надземная арматура
    IzolNadM2=0.0
    pc=p.filter(item__type__name='Кран',item__ust=UST_NAD)
    if pc:
        l.append(("Монтаж и покрытие 1 слоем эпоксидной грунт–эмали ИЗОЛЭП-mastic и 1 слоем эмали ПОЛИТОН-УР(УФ) надземной арматуры:",))
        for k in pc:
            l.append(('-'+k.item.name.encode('utf-8'),
                "шт"+CR+"кг"+CR+"м2",
                sf(k.numNad)+CR+
                kf(k.numNad*k.item.m)+CR+
                af(k.numNad*KranArea[k.item.dy])))
            IzolNadM2=IzolNadM2+k.numNad*KranArea[k.item.dy]

    # надземные трубопроводы без изоляции
    # в цикле перебираем все ДУ для труб и для фасонных частей начиная с максимальных
    # и выбираем подземные трубы и фасонину без изоляции для каждого Ду
    flagTitle=False
    for dy in reversed(sorted(DYD.keys())):
        flagTruba=False
        summ=0.0
        listIdx=0
        for d in DYD[dy]:
            pc=p.filter(item__type__name='Труба', item__d=d, item__zav_is=False,numNad__gt=0)
            if pc:
                flagTruba=True
                if not flagTitle:
                    flagTitle=True
                    l.append(('Монтаж и покрытие 1 слоем эпоксидной грунт–эмали ИЗОЛЭП-mastic и 1 слоем эмали ПОЛИТОН-УР(УФ) надземных',))
                    l.append(['трубопроводов Ду '+ str(dy) + ' мм в том числе:','м'+CR+'м2'])
                else:
                    l.append(['то же, трубопроводов Ду '+ str(dy) + ' мм в том числе:','м'+CR+'м2'])
                listIdx=len(l)
                for k in pc:
                    l.append(('-'+k.item.name.encode('utf-8')+' '+k.item.tu.name.encode('utf-8'),
                        'м'+CR+'кг',
                        mf(k.numNad)+CR+kf(k.numNad*k.item.m)))
                    summ+=k.numNad
        for d in DYD[dy]:
            pc=p.filter(item__type__name__in=GasList, item__d=d, item__zav_is=False, numNad__gt=0)
            if pc:
                if not flagTitle:
                    flagTitle=True
                    l.append(('Монтаж и покрытие 1 слоем эпоксидной грунт–эмали ИЗОЛЭП-mastic и 1 слоем эмали ПОЛИТОН-УР(УФ) надземных',))
                    l.append(['стальных фасонных частей Ду '+ str(dy) + ' мм в том числе:','м'+CR+'м2'])
                    listIdx=len(l)
                elif not flagTruba:
                    l.append(['то же, стальных фасонных частей Ду '+ str(dy) + ' мм в том числе:','м'+CR+'м2'])
                    listIdx=len(l)
                for k in pc:
                    l.append(('-'+k.item.name.encode('utf-8'),
                        'шт'+CR+'кг',
                        sf(k.numNad)+CR+kf(k.numNad*k.item.m)))
                    summ+=k.numNad*k.item.dlina/1000
        if summ>0:
            l[listIdx-1].append(mf(summ)+CR+af(AreaTrub(DYD[dy][0],summ)))
            IzolNadM2=IzolNadM2+AreaTrub(DYD[dy][0],summ)

    #Фланцы
    pc=p.filter(item__type__name='Фланец')
    if pc:
        l.append(("Монтаж и покрытие 1 слоем эпоксидной грунт–эмали ИЗОЛЭП-mastic и 1 слоем эмали ПОЛИТОН-УР(УФ) фланцев:",))
        for k in pc:
            FlanL,FlanD=Flanec7[k.item.dy]
            l.append(('-'+k.item.name.encode('utf-8'),
                "шт"+CR+"кг"+CR+"м2",
                sf(k.numAll)+CR+
                kf(k.numAll*k.item.m)+CR+
                af(k.numAll*AreaTrub(FlanD,FlanL/1000.0))))
            IzolNadM2=IzolNadM2+k.numAll*AreaTrub(FlanD,FlanL/1000.0)

    #монтаж монтажных изделий без заводской изоляцией
    pc=p.filter(item__type__name='Монтажные',item__zav_is=False)
    if pc:
        l.append(("Монтаж:",))
        for k in pc:
            l.append(('-'+k.item.name.encode('utf-8'), "шт"+CR+"кг", sf(k.numAll)+CR+kf(k.numAll*k.item.m)))

    #Предварительный подогрев стыков при сварке на трассе газопровода
    #принимаем что все стыки деталей идут с трубой (если это не так, то лишние стыки надо руками удалить из общего числа)
    #принимаем стыки кранов по ДУ - толщину уточнить после
    alls={}

    pc=p.filter(item__type__name='Тройник')
    if pc:
        for k in pc:
           dinc(alls,str(k.item.d)+"QQQ"+str(k.item.s),2*k.numAll)
           dinc(alls,str(k.item.d2)+"QQQ"+str(k.item.s2),k.numAll)

    pc=p.filter(item__type__name='Переход')
    if pc:
        for k in pc:
           dinc(alls,str(k.item.d)+"QQQ"+str(k.item.s),k.numAll)
           dinc(alls,str(k.item.d2)+"QQQ"+str(k.item.s2),k.numAll)

    pc=p.filter(item__type__name='Отвод')
    if pc:
        for k in pc:
           dinc(alls,str(k.item.d)+"QQQ"+str(k.item.s),2*k.numAll)

    pc=p.filter(item__type__name__in=['Заглушка','Днище'])
    if pc:
        for k in pc:
           dinc(alls,str(k.item.d)+"QQQ"+str(k.item.s),k.numAll)

    pc=p.filter(item__type__name='Кран')
    if pc:
        for k in pc:
           dinc(alls,"-QQQКран Ду%d, %dшт" % (k.item.dy, k.numAll), 2*k.numAll)

    pc=p.filter(item__type__name='Фланец')
    if pc:
        for k in pc:
           dinc(alls,"-QQQФланец Ду%d, %dшт" % (k.item.dy, k.numAll), k.numAll)

    l.append(("Предварительный подогрев стыков при сварке на трассе трубопровода:",))
    for k in sorted(alls.keys()):
        d, s = k.split("QQQ")
        l.append((d+"x"+s,"стык",sf(alls[k])))

    l.append(("Радиографический контроль качества сварных соединений импульсным рентгеновским аппаратом на трассе трубопровода:",))
    for k in sorted(alls.keys()):
        d, s = k.split("QQQ")
        l.append((d+"x"+s,"стык",sf(alls[k])))

    l.append(("Дополнительные затраты на обработку плёнки и расшифровку результатов для трубопровода:",))
    for k in sorted(alls.keys()):
        d, s = k.split("QQQ")
        if d=="-":
            l.append((s,"стык",str(alls[k])))
        else:
            l.append(("Ду"+str(GetDU(d)),"стык",sf(alls[k])))

    l.append(("Контроль качества сварных стыков дублирующим ультразвуковым методом:",))
    for k in sorted(alls.keys()):
        d, s = k.split("QQQ")
        l.append((d+"x"+s,"стык",sf(alls[k])))

    l.append(("Визуальный и измерительный контроль сварных соединений:",))
    for k in sorted(alls.keys()):
        d, s = k.split("QQQ")
        l.append((d+"x"+s,"стык",sf(alls[k])))

#Торцы не считаем. Выводим все трубы какие только есть в проекте
    pc=p.filter(item__type__name='Труба')
    if pc:
        l.append(("100 % ультразвуковой контроль всего периметра участка трубы на ширине не менее 40 мм от резаного торца с очисткой металлическими щётками и протиркой ацетоном:",))
        for k in pc:
            l.append(('%dx%.1f' % (k.item.d, k.item.s),'торец','ХХХ'))

    s=Spec.objects.get(pk=spec_id)
    l.append(('Предварительное испытание узла Ду ХХХ мм на прочность Рисп=1.1Рраб=%.2f МПа, в том числе:' % (1.1*s.Prab),
        'узел','ХХХ'))
    l.append(('-задвижки: Ду ХХХ мм','шт','ХХХ'))
    l.append(('-трубы: Ду ХХХ мм','м','ХХХ'))

    l.append(('То же: Ду ХХХ мм', 'узел','ХХХ'))
    l.append(('-задвижки: Ду ХХХ мм','шт','ХХХ'))
    l.append(('-трубы: Ду ХХХ мм','м','ХХХ'))

#Считаем общие длины всех труб, деталей и кранов по ДУ от больших к меньшим
    trubs=[]
    for dy in reversed(sorted(DYD.keys())):
        summ=0
        for d in DYD[dy]:
            pc=p.filter(item__type__name='Труба', item__d=d, isHydroIsp=True)
            for k in pc:
                summ+=k.numAll
            pc=p.filter(item__type__name__in=GasList, item__d=d, isHydroIsp=True)
            for k in pc:
                summ+=k.numAll*k.item.dlina/1000
        pc=p.filter(item__type__name='Кран', item__dy=dy, isHydroIsp=True)
        for k in pc:
            summ+=k.numAll*k.item.dlina/1000
        if summ>0:
            trubs.append((dy,summ))

    l.append(('Очистка полости трубопровода водой:',))
    for j in trubs:
        l.append(('Ду %d мм' % j[0],'м','%.1f' % j[1]))

#испытания на прочность
    s=Spec.objects.get(pk=spec_id)
    trays = s.NumTry.split()
    while len(trays)>1:
        koaff=float(trays.pop(0))
        l.append(('Предварительное гидравлическое испытание на прочность Рисп=%.2fРраб=%.2f МПа трубопровода:' % (koaff, koaff*s.Prab),))
        for j in trubs:
            l.append(('Ду %d мм' % j[0],'м','%.1f' % j[1]))

    koaff=float(trays.pop(0))
    l.append(('Гидравлическое испытание на прочность Рисп=%.1fРраб=%.1f МПа в течение 12 часов и герметичность Рисп=Рраб=%.1f МПа трубопровода в течение времени необходимого на осмотр, но не менее 12 часов:' % (koaff, koaff*s.Prab, s.Prab),))
    for j in trubs:
        l.append(('Ду %d мм' % j[0],'м','%.1f' % j[1]))

    l.append(('Осушка трубопровода:',))
    for j in trubs:
        l.append(('Ду %d мм' % j[0],'м','%.1f' % j[1]))

#Внутренний объем +15% на смачивание
    pc=p.filter(item__type__name__in=GasFullList, isHydroIsp=True)
    v=0.0
    for k in pc:
        if k.item.type.name.encode('utf-8')=='Труба':
            v+=ValTrub(k.item.d-2*k.item.s, k.numAll)
        elif k.item.type.name.encode('utf-8')=='Кран':
            v+=ValTrub(max(DYD[k.item.dy]), k.numAll*k.item.dlina/1000)
        elif k.item.type.name.encode('utf-8')=='Отвод':
            v+=ValTrub(k.item.d-2*k.item.s, k.numAll*k.item.dlina/1000)
        elif k.item.type.name.encode('utf-8')=='Переход':
            r=(k.item.d/2-k.item.s)/1000
            r1=(k.item.d2/2-k.item.s2)/1000
            v+=(r*r+r1*r1+r*r1)*pi*(k.item.dlina/1000)/3*k.numAll
        elif k.item.type.name.encode('utf-8')=='Тройник':
            v+=ValTrub(k.item.d-2*k.item.s, k.numAll*k.item.dlina/1000)
            v+=ValTrub(k.item.d2-2*k.item.s2, k.numAll*(k.item.visota-k.item.d/2)/1000)

    l.append(('Объём воды для гидроиспытаний трубопровода с учетом 15% на смачивание','м3',mf(v*1.15)))

    l.append(('Установка знака «Запрещается пользоваться открытым огнем и курить»','шт','2'))
    l.append(('То же, знак «Конденсат! Вход запрещен»','шт','2'))
    l.append(('То же, информационная табличка с указанием ЭО, филиала ЭО и телефона филиала ЭО ','шт','2'))

    l.append(('Монтаж и демонтаж днищ для испытаний:',))
    l.append(('-днище ХХХ','шт'+CR+'кг',))

    l.append(('Разработка котлована под ХХХ и обратная засыпка в равнинно-холмистой местности в грунтах ХХХ гр.','м3','ХХХ'))

#выводим суммарную площадь для покрытия краской/биурс
    l.append(('Итого под БИУРС','м2','%.1f' % IzolPodM2))
    l.append(('Итого под эмаль','м2','%.1f' % IzolNadM2))


#Врезка под давлением
    pc=p.filter(item__type__name='Врезка под давлением')
    if pc:
        l.append(("Врезка под давлением оборудованием фирмы T.D. Williamson",))

    document=Document("templateVR.docx")
    t1=document.tables[0]
    i=0
    for k in l:
        if k[0]:
            t1.cell(i,1).text=k[0].decode('utf-8')
        if len(k)>1:
            t1.cell(i,2).text=k[1].decode('utf-8')
        if len(k)>2:
            t1.cell(i,5).text=k[2].decode('utf-8')
        i=i+1
        t1.add_row()
    f=NamedTemporaryFile(delete=False)
    f.close()
    document.save(f.name)
    f=open(f.name,'rb')
    response = HttpResponse(f, content_type='text/html')
    response['Content-Disposition'] = 'attachment; filename="VR.docx"'
    return response

#найти Ду для данного Д
def GetDU(d):
    for k in DYD:
        if int(d) in DYD[k]:
            return k
    return 0

#если в словаре d есть значение с ключем k оно увеличивается на v иначе создается значение с ключем k и значением v
def dinc(d,k,v):
    if not k in d:
        d[k]=v
    else:
        d[k]+=v

def sf(i):
    return ShtukFormat.format(i)

def af(i):
    return AreaFormat.format(i)

def kf(i):
    return MassaFormat.format(i)

def mf(i):
    return MetrFormat.format(i)

# площадь наружной поверхности трубы м2. d - внешний диаметр трубы, мм; l - длина трубы, м
def AreaTrub(d,l):
    return l*pi*d/1000

# объем трубы м3. d - внутренний диаметр трубы, мм; l - длина трубы, м
def ValTrub(d,l):
    return pi*d*d*l/4/1000/1000

def tohtmlteg(s):
    data='<html><table>'
    for k in s:
        data+='<tr>'
        for j in k:
            data+=td(j)
        data+='</tr>'
    data+='</table></html>'
    return data

def tohtml(s):
    of=open("vr_file.html","w")
    of.write('<html><table>')
    for k in s:
        of.write('<tr>')
        for j in k:
            of.write(td(j))
        of.write('</tr>')
    of.write('</table></html>')
    of.close()

def td(s):
    return '<td>'+s+'</td>'

def tr(s):
    return '<tr>'+s+'</tr>'
